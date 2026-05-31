import streamlit as st
import pandas as pd
import numpy as np
import folium
from folium.plugins import HeatMap
from streamlit_folium import st_folium
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, RGBColor
from io import BytesIO
from datetime import datetime
import plotly.express as px
import plotly.graph_objects as go

# ─── Page Config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Urban Heat Island Analyzer",
    page_icon=None,
    layout="wide"
)

# ─── Theme Switcher Session State & Callback ────────────────────────────────────
if "theme" not in st.session_state:
    st.session_state["theme"] = "dark"  # Default to dark mode

def toggle_theme():
    if st.session_state["theme"] == "dark":
        st.session_state["theme"] = "light"
    else:
        st.session_state["theme"] = "dark"

# ─── CSS Variables Setup based on Active Theme ─────────────────────────────────
if st.session_state["theme"] == "dark":
    theme_variables = """
    :root {
        --bg:                 #0A0A0A;
        --surface:            #181818; /* Card Surface */
        --surface-secondary:  #141414; /* Main Surface */
        --surface-tertiary:   #1B1B1B; /* Sidebar Surface */
        --text-primary:       #F5F5F5;
        --text-secondary:     #CFCFCF;
        --text-muted:         #9A9A9A;
        --border:             #2D2D2D;
        --hover:              #222222;
        --accent-primary:     #4B5563;
        --accent-success:     #059669;
        --accent-warning:     #D97706;
        --accent-danger:      #DC2626;
        --accent-primary-hover: #5A6578;
        --header-shadow:      rgba(0, 0, 0, 0.8);
    }
    """
else:
    theme_variables = """
    :root {
        --bg:                 #F5F7FA;
        --surface:            #FFFFFF; /* Card Surface */
        --surface-secondary:  #EEF2F7; /* Main Surface */
        --surface-tertiary:   #F9FAFB; /* Sidebar Surface */
        --text-primary:       #111827;
        --text-secondary:     #4B5563;
        --text-muted:         #6B7280;
        --border:             #D9E0EA;
        --hover:              #EEF2F7;
        --accent-primary:     #2563EB;
        --accent-success:     #059669;
        --accent-warning:     #D97706;
        --accent-danger:      #DC2626;
        --accent-primary-hover: #1D4ED8;
        --header-shadow:      rgba(0, 0, 0, 0.05);
    }
    """

# ─── Styling Injection ──────────────────────────────────────────────────────────
st.markdown(f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Poppins:wght@400;500;600;700&display=swap');

{theme_variables}

/* ══════════════════════════════════════════════
   GLOBAL STYLES & TYPOGRAPHY
   ══════════════════════════════════════════════ */
* {{
    box-sizing: border-box;
}}

html, body, [data-testid="stAppViewContainer"], .main, h1, h2, h3, h4, h5, h6, p, li, label, input, select, textarea, table, td, th {{
    font-family: 'Poppins', sans-serif !important;
}}

/* Safeguard Streamlit's Material Icons/Symbols from font family overrides */
.notranslate,
[class*="material-symbols"],
[class*="MaterialSymbols"],
[data-testid="stHeader"] button span,
[data-testid="stSidebar"] button span,
[data-testid="stSidebarCollapsedControl"] span,
[data-testid="stUploadedFile"] span {{
    font-family: "Material Symbols Outlined", "Material Symbols Rounded", "Material Icons" !important;
}}

html, body, [data-testid="stAppViewContainer"], .main {{
    background-color: var(--bg) !important;
    color: var(--text-primary) !important;
}}

/* Centered horizontally layout with generous gutters */
.block-container {{
    max-width: 1400px !important;
    margin: 0 auto !important;
    padding-top: 100px !important; /* height of sticky top nav + spacing */
    padding-left: 2rem !important;
    padding-right: 2rem !important;
    padding-bottom: 3rem !important;
    background-color: var(--bg) !important;
}}

/* Typography Hierarchy */
h1 {{
    font-size: 32px !important;
    font-weight: 700 !important;
    color: var(--text-primary) !important;
    margin-top: 0 !important;
    margin-bottom: 1rem !important;
    letter-spacing: -0.02em !important;
}}

h2 {{
    font-size: 24px !important;
    font-weight: 600 !important;
    color: var(--text-primary) !important;
    margin-top: 1.5rem !important;
    margin-bottom: 0.75rem !important;
    letter-spacing: -0.01em !important;
}}

h3 {{
    font-size: 18px !important;
    font-weight: 600 !important;
    color: var(--text-primary) !important;
    margin-top: 1.25rem !important;
    margin-bottom: 0.5rem !important;
}}

p, li, span {{
    font-size: 14px !important;
    font-weight: 400 !important;
    color: var(--text-secondary) !important;
}}

/* ══════════════════════════════════════════════
   STREAMLIT CORE UI OVERRIDES
   ══════════════════════════════════════════════ */
/* Hide native Streamlit headers background/decoration/footer but keep the sidebar control button */
[data-testid="stHeader"] {{
    background-color: transparent !important;
    border: none !important;
    box-shadow: none !important;
    height: 0px !important;
    pointer-events: none !important;
}}
/* Hide only the header toolbar menu (three-dots menu, status, etc.) */
[data-testid="stHeader"] [data-testid="stHeaderToolbar"] {{
    display: none !important;
}}

/* Ensure the sidebar collapse control button is styled beautifully as a floating menu toggle */
[data-testid="stHeader"] [data-testid="stSidebarCollapsedControl"] {{
    position: fixed !important;
    top: 86px !important;
    left: 16px !important;
    z-index: 1002 !important;
    background-color: var(--surface) !important;
    border: 1px solid var(--border) !important;
    border-radius: 8px !important;
    height: 40px !important;
    width: 40px !important;
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
    box-shadow: 0 2px 8px rgba(0, 0, 0, 0.1) !important;
    transition: background-color 0.2s ease, border-color 0.2s ease !important;
    pointer-events: auto !important;
}}
[data-testid="stHeader"] [data-testid="stSidebarCollapsedControl"] button {{
    background-color: transparent !important;
    border: none !important;
    color: var(--text-primary) !important;
}}
[data-testid="stHeader"] [data-testid="stSidebarCollapsedControl"] button:hover {{
    color: var(--accent-primary) !important;
}}

[data-testid="stDecoration"] {{
    display: none !important;
}}
footer {{
    display: none !important;
}}

/* Sidebar Custom Styling */
[data-testid="stSidebar"] {{
    background-color: var(--surface-tertiary) !important;
    border-right: 1px solid var(--border) !important;
    width: 280px !important;
    min-width: 280px !important;
    max-width: 280px !important;
    padding-top: 72px !important; /* push down below the top navigation bar */
    z-index: 99 !important;
}}

[data-testid="stSidebar"] .block-container {{
    padding: 1.5rem !important;
    max-width: 100% !important;
}}

[data-testid="stSidebar"] div[data-testid="stForm"] {{
    border: none !important;
    padding: 0 !important;
}}

/* Text Input & Select Box styling overrides */
div[data-baseweb="input"],
div[data-baseweb="base-input"],
input {{
    background-color: var(--surface) !important;
    color: var(--text-primary) !important;
    border-color: var(--border) !important;
}}
div[data-baseweb="input"] {{
    border: 1px solid var(--border) !important;
    border-radius: 10px !important;
    height: 42px !important;
    transition: border-color 0.2s ease, box-shadow 0.2s ease !important;
}}
div[data-baseweb="input"]:focus-within {{
    border-color: var(--accent-primary) !important;
    box-shadow: 0 0 0 1px var(--accent-primary) !important;
}}
input {{
    font-size: 14px !important;
}}
input::placeholder {{
    color: var(--text-muted) !important;
    opacity: 1 !important;
}}

/* Selectbox */
div[data-baseweb="select"],
div[data-baseweb="select"] > div {{
    background-color: var(--surface) !important;
    color: var(--text-primary) !important;
    border-color: var(--border) !important;
}}
div[data-baseweb="select"] {{
    border: 1px solid var(--border) !important;
    border-radius: 10px !important;
    height: 42px !important;
}}

/* Tabs Override */
div[data-baseweb="tab-list"] {{
    background-color: transparent !important;
    border-bottom: 1px solid var(--border) !important;
    gap: 8px !important;
    margin-bottom: 1.5rem !important;
}}
button[data-baseweb="tab"] {{
    font-family: 'Poppins', sans-serif !important;
    font-size: 14px !important;
    font-weight: 500 !important;
    color: var(--text-muted) !important;
    background-color: transparent !important;
    border: none !important;
    padding: 10px 18px !important;
    border-bottom: 2px solid transparent !important;
    transition: all 0.2s ease !important;
}}
button[data-baseweb="tab"]:hover {{
    color: var(--text-secondary) !important;
}}
button[data-baseweb="tab"][aria-selected="true"] {{
    color: var(--accent-primary) !important;
    border-bottom: 2px solid var(--accent-primary) !important;
    font-weight: 600 !important;
}}

/* Native Streamlit Buttons override */
div.stButton > button {{
    height: 42px !important;
    border-radius: 10px !important;
    font-size: 14px !important;
    font-weight: 500 !important;
    padding: 0 20px !important;
    transition: all 0.25s ease !important;
    font-family: 'Poppins', sans-serif !important;
}}
/* Primary Button */
div.stButton > button[kind="primary"] {{
    background-color: var(--accent-primary) !important;
    color: #FFFFFF !important;
    border: none !important;
}}
div.stButton > button[kind="primary"]:hover,
div.stButton > button[kind="primary"]:focus,
div.stButton > button[kind="primary"]:active {{
    background-color: var(--accent-primary-hover) !important;
    color: #FFFFFF !important;
    border: none !important;
    transform: translateY(-1px) !important;
    box-shadow: 0 4px 12px rgba(0, 0, 0, 0.15) !important;
}}
/* Secondary Button */
div.stButton > button[kind="secondary"] {{
    background-color: var(--surface) !important;
    color: var(--text-primary) !important;
    border: 1px solid var(--border) !important;
}}
div.stButton > button[kind="secondary"]:hover,
div.stButton > button[kind="secondary"]:focus,
div.stButton > button[kind="secondary"]:active {{
    background-color: var(--hover) !important;
    border-color: var(--text-muted) !important;
    color: var(--text-primary) !important;
    transform: translateY(-1px) !important;
}}

/* Download Button styling */
div.stDownloadButton > button {{
    height: 42px !important;
    border-radius: 10px !important;
    font-size: 14px !important;
    font-weight: 500 !important;
    padding: 0 20px !important;
    background-color: var(--accent-success) !important;
    color: #FFFFFF !important;
    border: none !important;
    transition: all 0.25s ease !important;
}}
div.stDownloadButton > button:hover,
div.stDownloadButton > button:focus,
div.stDownloadButton > button:active {{
    background-color: #047857 !important;
    color: #FFFFFF !important;
    transform: translateY(-1px) !important;
    box-shadow: 0 4px 12px rgba(0, 0, 0, 0.15) !important;
}}

/* File Uploader styling overrides */
[data-testid="stFileUploader"] {{
    background-color: var(--surface) !important;
    border: 1px dashed var(--border) !important;
    border-radius: 10px !important;
    padding: 0.85rem !important;
    transition: border-color 0.2s ease !important;
}}
[data-testid="stFileUploader"]:hover {{
    border-color: var(--accent-primary) !important;
}}
[data-testid="stFileUploader"] section {{
    padding: 0 !important;
    background: transparent !important;
}}
[data-testid="stFileUploader"] [data-testid="stUploadedFile"] {{
    background-color: var(--surface-secondary) !important;
    border: 1px solid var(--border) !important;
    border-radius: 8px !important;
}}
[data-testid="stFileUploader"] [data-testid="stUploadedFile"] * {{
    color: var(--text-primary) !important;
}}

/* ══════════════════════════════════════════════
   TOP NAVIGATION BAR
   ══════════════════════════════════════════════ */
.top-nav {{
    position: fixed;
    top: 0;
    left: 0;
    right: 0;
    height: 72px;
    background-color: var(--surface) !important;
    border-bottom: 1px solid var(--border) !important;
    display: flex;
    align-items: center;
    justify-content: space-between;
    padding: 0 2rem;
    z-index: 1000;
    box-shadow: 0 2px 10px var(--header-shadow);
}}
.top-nav-left {{
    display: flex;
    align-items: center;
    gap: 1.25rem;
}}
.logo-mark {{
    width: 40px;
    height: 40px;
    background-color: var(--accent-primary) !important;
    border-radius: 6px;
    display: flex;
    align-items: center;
    justify-content: center;
    font-size: 0.85rem;
    font-weight: 700;
    color: #FFFFFF !important;
    letter-spacing: 0.05em;
    text-transform: uppercase;
}}
.nav-brand {{
    display: flex;
    flex-direction: column;
}}
.nav-title {{
    font-size: 1.05rem;
    font-weight: 700;
    color: var(--text-primary) !important;
    line-height: 1.2;
}}
.nav-subtitle {{
    font-size: 0.72rem;
    color: var(--text-muted) !important;
    font-weight: 400;
}}
.nav-divider {{
    width: 1px;
    height: 32px;
    background-color: var(--border) !important;
}}
.nav-module {{
    font-size: 0.75rem;
    font-weight: 600;
    color: var(--text-secondary) !important;
    text-transform: uppercase;
    letter-spacing: 0.08em;
    background-color: var(--surface-secondary) !important;
    padding: 4px 10px;
    border-radius: 4px;
    border: 1px solid var(--border) !important;
}}
.top-nav-right {{
    display: flex;
    align-items: center;
    gap: 1.5rem;
    margin-right: 150px; /* leaves room for absolute-positioned theme button */
}}
.system-status {{
    display: flex;
    align-items: center;
    gap: 8px;
}}
.status-dot {{
    width: 8px;
    height: 8px;
    background-color: var(--accent-success) !important;
    border-radius: 50%;
    display: inline-block;
    box-shadow: 0 0 8px var(--accent-success);
}}
.status-text {{
    font-size: 0.75rem;
    font-weight: 500;
    color: var(--text-secondary) !important;
}}

/* Absolute-positioned Theme Toggler inside fixed top nav */
.theme-btn-wrapper {{
    position: fixed;
    top: 15px;
    right: 2rem;
    z-index: 1001;
}}
.theme-btn-wrapper div.stButton > button[kind="secondary"] {{
    height: 42px !important;
    border-radius: 10px !important;
    background-color: var(--surface-secondary) !important;
    border: 1px solid var(--border) !important;
    color: var(--text-secondary) !important;
    font-size: 12px !important;
    font-weight: 600 !important;
    text-transform: uppercase !important;
    letter-spacing: 0.05em !important;
    padding: 0 16px !important;
}}
.theme-btn-wrapper div.stButton > button[kind="secondary"]:hover,
.theme-btn-wrapper div.stButton > button[kind="secondary"]:focus,
.theme-btn-wrapper div.stButton > button[kind="secondary"]:active {{
    background-color: var(--hover) !important;
    border-color: var(--accent-primary) !important;
    color: var(--text-primary) !important;
    transform: translateY(-1px) !important;
}}

/* ══════════════════════════════════════════════
   SIDEBAR COMPONENT STYLES
   ══════════════════════════════════════════════ */
.sb-section {{
    font-size: 11px !important;
    font-weight: 700 !important;
    color: var(--text-muted) !important;
    text-transform: uppercase !important;
    letter-spacing: 0.12em !important;
    margin-top: 1.5rem !important;
    margin-bottom: 0.5rem !important;
    padding-bottom: 4px !important;
    border-bottom: 1px solid var(--border) !important;
}}
.sb-nav {{
    display: flex;
    flex-direction: column;
    gap: 4px;
    margin-bottom: 0.5rem;
}}
.sb-nav-item {{
    display: flex;
    align-items: center;
    gap: 10px;
    padding: 10px 12px;
    border-radius: 8px;
    font-size: 13px;
    font-weight: 500;
    color: var(--text-secondary);
    background-color: transparent;
    transition: all 0.2s ease;
    cursor: default;
    border: 1px solid transparent;
}}
.sb-nav-item.active {{
    color: var(--text-primary);
    background-color: var(--surface);
    border: 1px solid var(--border);
}}
.sb-nav-dot {{
    width: 6px;
    height: 6px;
    border-radius: 50%;
    background-color: var(--text-muted);
}}
.sb-nav-item.active .sb-nav-dot {{
    background-color: var(--accent-primary);
    box-shadow: 0 0 6px var(--accent-primary);
}}
.sb-meta {{
    font-size: 12px !important;
    color: var(--text-muted) !important;
    line-height: 1.8 !important;
    margin-top: 1.5rem !important;
    padding-top: 1rem !important;
    border-top: 1px solid var(--border) !important;
}}

/* ══════════════════════════════════════════════
   METRIC CARDS
   ══════════════════════════════════════════════ */
.metric-row {{
    display: grid;
    grid-template-columns: repeat(4, 1fr);
    gap: 16px;
    margin-bottom: 1.8rem;
}}

.mcard {{
    background-color: var(--surface) !important;
    border: 1px solid var(--border) !important;
    border-radius: 12px !important;
    padding: 1.25rem !important;
    box-shadow: 0 1px 3px rgba(0, 0, 0, 0.05) !important;
    transition: transform 0.2s ease, box-shadow 0.2s ease !important;
    display: flex;
    flex-direction: column;
    gap: 10px;
}}
.mcard:hover {{
    transform: translateY(-2px) !important;
    box-shadow: 0 4px 12px rgba(0, 0, 0, 0.08) !important;
}}
.mcard.blue {{ border-left: 4px solid var(--accent-primary) !important; }}
.mcard.orange {{ border-left: 4px solid var(--accent-warning) !important; }}
.mcard.red {{ border-left: 4px solid var(--accent-danger) !important; }}
.mcard.green {{ border-left: 4px solid var(--accent-success) !important; }}

.mcard-header {{
    display: flex;
    justify-content: space-between;
    align-items: center;
}}
.mcard-label {{
    font-size: 11px !important;
    font-weight: 600 !important;
    text-transform: uppercase !important;
    letter-spacing: 0.08em !important;
    color: var(--text-muted) !important;
}}
.mcard-value {{
    font-size: 28px !important;
    font-weight: 700 !important;
    color: var(--text-primary) !important;
    letter-spacing: -0.02em !important;
    line-height: 1.1 !important;
}}
.mcard-metadata {{
    font-size: 12px !important;
    font-weight: 400 !important;
    color: var(--text-secondary) !important;
}}
.mcard-indicator {{
    font-size: 10px !important;
    font-weight: 600 !important;
    text-transform: uppercase !important;
    letter-spacing: 0.05em !important;
    padding: 2px 8px;
    border-radius: 6px;
    border: 1px solid transparent !important;
}}
.mcard-indicator.blue {{
    background-color: rgba(37, 99, 235, 0.08) !important;
    color: var(--accent-primary) !important;
    border-color: rgba(37, 99, 235, 0.15) !important;
}}
.mcard-indicator.orange {{
    background-color: rgba(217, 119, 6, 0.08) !important;
    color: var(--accent-warning) !important;
    border-color: rgba(217, 119, 6, 0.15) !important;
}}
.mcard-indicator.red {{
    background-color: rgba(220, 38, 38, 0.08) !important;
    color: var(--accent-danger) !important;
    border-color: rgba(220, 38, 38, 0.15) !important;
}}
.mcard-indicator.green {{
    background-color: rgba(5, 150, 105, 0.08) !important;
    color: var(--accent-success) !important;
    border-color: rgba(5, 150, 105, 0.15) !important;
}}

/* ══════════════════════════════════════════════
   CHART & MAP CONTAINERS
   ══════════════════════════════════════════════ */
[data-testid="stPlotlyChart"] {{
    background: var(--surface) !important;
    border: 1px solid var(--border) !important;
    border-radius: 12px !important;
    padding: 1.25rem !important;
    box-shadow: 0 1px 3px rgba(0, 0, 0, 0.05) !important;
    transition: transform 0.2s ease, box-shadow 0.2s ease !important;
}}
[data-testid="stPlotlyChart"]:hover {{
    transform: translateY(-2px) !important;
    box-shadow: 0 4px 12px rgba(0, 0, 0, 0.08) !important;
}}

.map-container {{
    background: var(--surface) !important;
    border: 1px solid var(--border) !important;
    border-radius: 12px !important;
    padding: 1rem !important;
    box-shadow: 0 1px 3px rgba(0, 0, 0, 0.05) !important;
    margin-bottom: 1.5rem !important;
    transition: transform 0.2s ease, box-shadow 0.2s ease !important;
}}
.map-container:hover {{
    transform: translateY(-2px) !important;
    box-shadow: 0 4px 12px rgba(0, 0, 0, 0.08) !important;
}}
.map-container iframe {{
    border-radius: 8px !important;
}}

/* Section Titles */
.section-title {{
    font-size: 12px !important;
    font-weight: 700 !important;
    color: var(--text-muted) !important;
    text-transform: uppercase !important;
    letter-spacing: 0.12em !important;
    padding-bottom: 8px !important;
    border-bottom: 1px solid var(--border) !important;
    margin-top: 1.5rem !important;
    margin-bottom: 1.25rem !important;
}}

/* ══════════════════════════════════════════════
   DATA TABLES
   ══════════════════════════════════════════════ */
.gov-table-container {{
    background-color: var(--surface) !important;
    border: 1px solid var(--border) !important;
    border-radius: 12px !important;
    padding: 0 !important;
    box-shadow: 0 1px 3px rgba(0, 0, 0, 0.05) !important;
    overflow-x: auto !important;
    margin-bottom: 1.5rem !important;
    transition: transform 0.2s ease, box-shadow 0.2s ease !important;
}}
.gov-table-container:hover {{
    transform: translateY(-2px) !important;
    box-shadow: 0 4px 12px rgba(0, 0, 0, 0.08) !important;
}}
table.gov-table {{
    width: 100% !important;
    border-collapse: collapse !important;
    font-family: 'Poppins', sans-serif !important;
    font-size: 14px !important;
    text-align: left !important;
}}
table.gov-table th {{
    background-color: var(--surface-secondary) !important;
    color: var(--text-primary) !important;
    font-weight: 600 !important;
    padding: 14px 20px !important;
    border-bottom: 1px solid var(--border) !important;
    font-size: 12px !important;
    text-transform: uppercase !important;
    letter-spacing: 0.05em !important;
    position: sticky !important;
    top: 0 !important;
    z-index: 10 !important;
}}
table.gov-table td {{
    padding: 14px 20px !important;
    border-bottom: 1px solid var(--border) !important;
    color: var(--text-secondary) !important;
}}
table.gov-table tbody tr:last-child td {{
    border-bottom: none !important;
}}
table.gov-table tbody tr:nth-child(even) {{
    background-color: var(--surface-secondary) !important;
}}
table.gov-table tbody tr:hover {{
    background-color: var(--hover) !important;
}}

/* ══════════════════════════════════════════════
   INFO & WARNING BANNERS
   ══════════════════════════════════════════════ */
.info-banner {{
    background-color: var(--surface-secondary) !important;
    border: 1px solid var(--border) !important;
    border-left: 4px solid var(--accent-primary) !important;
    border-radius: 8px !important;
    padding: 1.2rem 1.5rem !important;
    color: var(--text-secondary) !important;
    font-size: 13px !important;
    line-height: 1.5 !important;
    margin-bottom: 1.5rem !important;
}}
.info-banner strong {{
    color: var(--text-primary) !important;
    font-weight: 600 !important;
}}

/* ══════════════════════════════════════════════
   AI MITIGATION REPORT TEXT STYLING
   ══════════════════════════════════════════════ */
.report-container p, .report-container li {{
    font-size: 14px !important;
    line-height: 1.6 !important;
    color: var(--text-secondary) !important;
}}
.report-container h1, .report-container h2, .report-container h3 {{
    color: var(--text-primary) !important;
    font-family: 'Poppins', sans-serif !important;
}}
.report-container h1 {{
    font-size: 24px !important;
    font-weight: 700 !important;
    border-bottom: 1px solid var(--border) !important;
    padding-bottom: 8px !important;
    margin-top: 1.8rem !important;
}}
.report-container h2 {{
    font-size: 18px !important;
    font-weight: 600 !important;
    margin-top: 1.5rem !important;
}}

/* ══════════════════════════════════════════════
   WELCOME / EMPTY STATES
   ══════════════════════════════════════════════ */
.welcome-screen {{
    display: flex;
    flex-direction: column;
    gap: 2rem;
    margin-top: 0.5rem;
    max-width: 1200px;
}}
.welcome-hero {{
    border-bottom: 1px solid var(--border) !important;
    padding-bottom: 1.5rem !important;
}}
.welcome-title {{
    font-size: 2.2rem !important;
    font-weight: 700 !important;
    color: var(--text-primary) !important;
    margin-bottom: 0.5rem !important;
    letter-spacing: -0.03em !important;
}}
.welcome-subtitle {{
    font-size: 1.05rem !important;
    font-weight: 400 !important;
    color: var(--text-secondary) !important;
    line-height: 1.4 !important;
}}
.welcome-grid {{
    display: grid;
    grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
    gap: 1.5rem;
}}
.welcome-card {{
    background-color: var(--surface) !important;
    border: 1px solid var(--border) !important;
    border-radius: 12px !important;
    padding: 1.5rem !important;
    box-shadow: 0 1px 3px rgba(0, 0, 0, 0.05) !important;
    display: flex;
    flex-direction: column;
    gap: 10px;
}}
.welcome-card-header {{
    display: flex;
    flex-direction: column;
    gap: 6px;
}}
.welcome-card-tag {{
    font-size: 9px !important;
    font-weight: 600 !important;
    text-transform: uppercase !important;
    letter-spacing: 0.08em !important;
    color: var(--accent-primary) !important;
    background-color: rgba(37, 99, 235, 0.08) !important;
    padding: 2px 6px !important;
    border-radius: 4px !important;
    width: fit-content;
}}
.welcome-card h3 {{
    font-size: 16px !important;
    font-weight: 600 !important;
    color: var(--text-primary) !important;
    margin: 0 !important;
}}
.welcome-card-desc {{
    font-size: 13px !important;
    line-height: 1.6 !important;
    color: var(--text-secondary) !important;
    margin: 0 !important;
}}
.welcome-cta-card {{
    display: flex;
    align-items: center;
    gap: 1.25rem;
    background-color: var(--surface-secondary) !important;
    border: 1px solid var(--border) !important;
    border-radius: 12px !important;
    padding: 1.25rem 1.5rem !important;
}}
.cta-status-indicator {{
    width: 12px;
    height: 12px;
    background-color: var(--accent-primary) !important;
    border-radius: 50%;
    animation: cta-pulse 2.5s infinite ease-in-out;
}}
@keyframes cta-pulse {{
    0% {{ transform: scale(0.9); opacity: 0.5; }}
    50% {{ transform: scale(1.1); opacity: 1; }}
    100% {{ transform: scale(0.9); opacity: 0.5; }}
}}
.cta-content h4 {{
    font-size: 14px !important;
    font-weight: 600 !important;
    color: var(--text-primary) !important;
    margin: 0 0 2px 0 !important;
}}
.cta-content p {{
    font-size: 13px !important;
    color: var(--text-secondary) !important;
    margin: 0 !important;
}}
</style>
""", unsafe_allow_html=True)

# ─── Top Header Navigation Layout ──────────────────────────────────────────────
st.markdown("""
<div class="top-nav">
    <div class="top-nav-left">
        <div class="logo-mark">UHI</div>
        <div class="nav-brand">
            <span class="nav-title">Urban Heat Island Analyzer</span>
            <span class="nav-subtitle">Satellite-Based Land Surface Temperature Analysis Platform</span>
        </div>
        <div class="nav-divider"></div>
        <div class="nav-module">Climate Monitoring</div>
    </div>
    <div class="top-nav-right">
        <div class="system-status">
            <span class="status-dot"></span>
            <span class="status-text">NASA ECOSTRESS • Active</span>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)

# ─── Top Header Switcher Render ───────────────────────────────────────────────
st.markdown('<div class="theme-btn-wrapper">', unsafe_allow_html=True)
if st.session_state["theme"] == "dark":
    st.button("Light Mode", on_click=toggle_theme, key="theme_toggle_btn")
else:
    st.button("Dark Mode", on_click=toggle_theme, key="theme_toggle_btn")
st.markdown('</div>', unsafe_allow_html=True)

# ─── Sidebar Render ────────────────────────────────────────────────────────────
# Navigation Section
st.sidebar.markdown('<div class="sb-section">Navigation</div>', unsafe_allow_html=True)
st.sidebar.markdown("""
<div class="sb-nav">
    <div class="sb-nav-item active">
        <span class="sb-nav-dot"></span>
        LST Observatory Dashboard
    </div>
</div>
""", unsafe_allow_html=True)

# Filters Section
st.sidebar.markdown('<div class="sb-section">Filters</div>', unsafe_allow_html=True)
ward_name = st.sidebar.text_input("Ward Name", value="Hubballi-Dharwad", label_visibility="collapsed")

# Upload Section
st.sidebar.markdown('<div class="sb-section">Upload Section</div>', unsafe_allow_html=True)
uploaded_file = st.sidebar.file_uploader("Upload NASA LST CSV", type=["csv"], label_visibility="collapsed")

# Settings Section
st.sidebar.markdown('<div class="sb-section">Settings</div>', unsafe_allow_html=True)
gemini_key = st.sidebar.text_input("Gemini API Key", type="password", label_visibility="collapsed",
                                    placeholder="Enter Gemini API key…")

# Metadata Section
st.sidebar.markdown("""
<div class="sb-meta">
    Analysis Period: April – June 2023<br>
    Coordinate System: WGS84<br>
    Temperature Unit: Degrees Celsius<br>
    Cloud Filter: Clear pixels only
</div>
""", unsafe_allow_html=True)


# ─── Data Loading ──────────────────────────────────────────────────────────────
def load_data(file):
    df = pd.read_csv(file)
    df = df.rename(columns={"ECO_L2T_LSTE_002_LST": "LST_K"})
    df["LST_C"] = df["LST_K"] - 273.15
    df["Date"] = pd.to_datetime(df["Date"])
    df = df[df["LST_C"] > 0]  # Remove invalid readings
    df = df[df["ECO_L2T_LSTE_002_cloud_Cloud_Mask_Description"] == "Clear"]
    return df

# ─── Hotspot Detection ─────────────────────────────────────────────────────────
def detect_hotspots(df):
    mean_t = df["LST_C"].mean()
    std_t = df["LST_C"].std()

    def classify(t):
        if t > mean_t + 2 * std_t:
            return "Critical"
        elif t > mean_t + std_t:
            return "High"
        elif t > mean_t:
            return "Moderate"
        else:
            return "Normal"

    df["Zone"] = df["LST_C"].apply(classify)
    stats = {
        "mean": mean_t,
        "std": std_t,
        "max": df["LST_C"].max(),
        "min": df["LST_C"].min(),
        "critical_pct": len(df[df["Zone"] == "Critical"]) / len(df) * 100,
        "high_pct": len(df[df["Zone"] == "High"]) / len(df) * 100,
        "moderate_pct": len(df[df["Zone"] == "Moderate"]) / len(df) * 100,
    }
    return df, stats

# ─── Heatmap ───────────────────────────────────────────────────────────────────
def create_heatmap(df):
    center = [df["Latitude"].mean(), df["Longitude"].mean()]
    m = folium.Map(location=center, zoom_start=12, tiles="OpenStreetMap")

    heat_data = [[row["Latitude"], row["Longitude"], row["LST_C"]]
                 for _, row in df.iterrows()]
    HeatMap(heat_data, radius=30, blur=20,
            gradient={"0.4": "blue", "0.6": "yellow", "0.8": "orange", "1.0": "red"}
            ).add_to(m)

    zone_colors = {"Critical": "red", "High": "orange", "Moderate": "yellow", "Normal": "green"}
    for _, row in df.drop_duplicates("Category").iterrows():
        avg_temp = df[df["Category"] == row["Category"]]["LST_C"].mean()
        zone = df[df["Category"] == row["Category"]]["Zone"].mode()[0]
        folium.CircleMarker(
            location=[row["Latitude"], row["Longitude"]],
            radius=12,
            color=zone_colors.get(zone, "blue"),
            fill=True,
            popup=f"{row['Category']}<br>Avg LST: {avg_temp:.1f}°C<br>Zone: {zone}"
        ).add_to(m)

    return m

# ─── Factor Analysis ───────────────────────────────────────────────────────────
def analyse_factors(df):
    factors = {
        "Hubballi": {"buildup_pct": 72, "green_pct": 12, "impervious_pct": 68, "water_pct": 3},
        "Dharwad":  {"buildup_pct": 55, "green_pct": 28, "impervious_pct": 52, "water_pct": 5},
        "Hubballi-Central": {"buildup_pct": 85, "green_pct": 6,  "impervious_pct": 80, "water_pct": 1},
        "Dharwad-Central":  {"buildup_pct": 60, "green_pct": 22, "impervious_pct": 58, "water_pct": 4},
    }
    avg_temp_by_loc = df.groupby("Category")["LST_C"].mean().to_dict()
    return factors, avg_temp_by_loc

# ─── Gemini Report Generation ──────────────────────────────────────────────────
def generate_report(api_key, ward, stats, factors, avg_temps):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel("gemini-2.5-flash")

    factor_summary = "\n".join([
        f"- {loc}: Built-up {f['buildup_pct']}%, Green cover {f['green_pct']}%, "
        f"Avg LST {avg_temps.get(loc, 0):.1f}°C"
        for loc, f in factors.items()
    ])

    prompt = f"""
You are an expert urban planner and climate scientist specializing in UHI mitigation for Indian cities.

Generate a detailed, professional UHI Mitigation Strategy Report for {ward}, Karnataka, India.

## Data from NASA ECOSTRESS Satellite (April-June 2023):
- Mean Land Surface Temperature: {stats['mean']:.2f}°C
- Maximum LST recorded: {stats['max']:.2f}°C
- Minimum LST recorded: {stats['min']:.2f}°C
- Critical hotspot zones: {stats['critical_pct']:.1f}% of area
- High heat zones: {stats['high_pct']:.1f}% of area

## Location-wise Analysis:
{factor_summary}

Generate a report with these exact sections:
1. Executive Summary
2. UHI Hotspot Analysis (reference the actual temperature data)
3. Contributing Factors Analysis (built-up density, impervious surfaces, lack of green cover)
4. Tree Plantation Recommendations (include specific native Karnataka tree species like Neem, Peepal, Rain Tree, Tamarind - with quantities and locations)
5. Reflective Surface Recommendations (cool roofs, permeable pavements - with specific wards)
6. Water Body & Blue Infrastructure Recommendations
7. Implementation Priority Matrix (High/Medium/Low priority actions)
8. Expected Temperature Reduction (quantified estimates)
9. Conclusion

Be specific, data-driven, and practical for a Karnataka municipal context.
"""

    response = model.generate_content(prompt)
    return response.text

# ─── DOCX Export ───────────────────────────────────────────────────────────────
def create_docx(report_text, ward, stats):
    doc = Document()

    title = doc.add_heading(f"Urban Heat Island Mitigation Report", 0)
    title.alignment = 1
    doc.add_heading(f"{ward}, Karnataka, India", 1).alignment = 1
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')} | Data Source: NASA ECOSTRESS ECO_L2T_LSTE.002")

    doc.add_heading("Key Temperature Statistics", 2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"
    stats_data = [
        ("Mean LST", f"{stats['mean']:.2f}°C"),
        ("Maximum LST", f"{stats['max']:.2f}°C"),
        ("Minimum LST", f"{stats['min']:.2f}°C"),
        ("Critical Hotspot Area", f"{stats['critical_pct']:.1f}%"),
    ]
    for metric, value in stats_data:
        row = table.add_row().cells
        row[0].text = metric
        row[1].text = value

    doc.add_paragraph("")

    for line in report_text.split("\n"):
        if line.startswith("## ") or line.startswith("# "):
            doc.add_heading(line.replace("## ", "").replace("# ", ""), 2)
        elif line.startswith("### "):
            doc.add_heading(line.replace("### ", ""), 3)
        elif line.strip():
            doc.add_paragraph(line)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ─── Custom HTML Table Generator with Heat Gradient Column ─────────────────────
def get_heat_color(val, min_val, max_val, is_dark):
    if max_val == min_val:
        norm = 0.5
    else:
        norm = max(0.0, min(1.0, (val - min_val) / (max_val - min_val)))
    
    # Adaptive base heat colors (Success green -> Warning orange/yellow -> Danger red)
    c_green = (5, 150, 105)
    c_yellow = (217, 119, 6)
    c_red = (220, 38, 38)
    alpha = 0.22 if is_dark else 0.12
    
    if norm < 0.5:
        # Scale between Green (0.0) and Yellow (0.5)
        t = norm * 2
        r = int(c_green[0] + (c_yellow[0] - c_green[0]) * t)
        g = int(c_green[1] + (c_yellow[1] - c_green[1]) * t)
        b = int(c_green[2] + (c_yellow[2] - c_green[2]) * t)
    else:
        # Scale between Yellow (0.5) and Red (1.0)
        t = (norm - 0.5) * 2
        r = int(c_yellow[0] + (c_red[0] - c_yellow[0]) * t)
        g = int(c_yellow[1] + (c_red[1] - c_yellow[1]) * t)
        b = int(c_yellow[2] + (c_red[2] - c_yellow[2]) * t)
        
    return f"rgba({r}, {g}, {b}, {alpha})"

def generate_html_table(df, is_dark):
    min_mean = df["mean"].min()
    max_mean = df["mean"].max()
    
    html = '<div class="gov-table-container"><table class="gov-table">'
    html += '<thead><tr><th>Location</th><th>Mean Temp (°C)</th><th>Max Temp (°C)</th><th>Min Temp (°C)</th></tr></thead>'
    html += '<tbody>'
    
    for loc, row in df.iterrows():
        mean_val = row["mean"]
        max_val = row["max"]
        min_val = row["min"]
        
        bg_color = get_heat_color(mean_val, min_mean, max_mean, is_dark)
        
        html += '<tr>'
        html += f'<td><strong>{loc}</strong></td>'
        html += f'<td style="background-color: {bg_color}; font-weight: 700; color: var(--text-primary);">{mean_val:.2f}</td>'
        html += f'<td>{max_val:.2f}</td>'
        html += f'<td>{min_val:.2f}</td>'
        html += '</tr>'
        
    html += '</tbody></table></div>'
    return html

# ─── Main App ──────────────────────────────────────────────────────────────────
if uploaded_file:
    df = load_data(uploaded_file)
    df, stats = detect_hotspots(df)
    factors, avg_temps = analyse_factors(df)

    # ── Overview Metrics ──
    st.markdown('<div class="section-title">Overview Metrics</div>', unsafe_allow_html=True)

    # Metrics Row
    st.markdown(f"""
<div class="metric-row">
    <div class="mcard blue">
        <div class="mcard-header">
            <span class="mcard-label">Mean Surface Temp</span>
            <span class="mcard-indicator blue">Baseline</span>
        </div>
        <div class="mcard-value">{stats['mean']:.1f}°C</div>
        <div class="mcard-metadata">Dataset average (clear-sky)</div>
    </div>
    <div class="mcard orange">
        <div class="mcard-header">
            <span class="mcard-label">Maximum LST Recorded</span>
            <span class="mcard-indicator orange">High</span>
        </div>
        <div class="mcard-value">{stats['max']:.1f}°C</div>
        <div class="mcard-metadata">Peak surface temperature</div>
    </div>
    <div class="mcard red">
        <div class="mcard-header">
            <span class="mcard-label">Critical Hotspot Zones</span>
            <span class="mcard-indicator red">Critical</span>
        </div>
        <div class="mcard-value">{stats['critical_pct']:.1f}%</div>
        <div class="mcard-metadata">Area above mean + 2σ</div>
    </div>
    <div class="mcard green">
        <div class="mcard-header">
            <span class="mcard-label">Total Observations</span>
            <span class="mcard-indicator green">Valid</span>
        </div>
        <div class="mcard-value">{len(df):,}</div>
        <div class="mcard-metadata">Cloud-cleared data points</div>
    </div>
</div>
""", unsafe_allow_html=True)

    # ── Tabs ──
    tab1, tab2, tab3 = st.tabs(["Spatial Heatmap", "Statistical Analysis", "Mitigation Report"])

    with tab1:
        st.markdown('<div class="section-title">Land Surface Temperature — Spatial Distribution</div>', unsafe_allow_html=True)
        m = create_heatmap(df)
        st.markdown('<div class="map-container">', unsafe_allow_html=True)
        st_folium(m, width="100%", height=520)
        st.markdown('</div>', unsafe_allow_html=True)

    with tab2:
        st.markdown('<div class="section-title">Average Temperature by Location</div>', unsafe_allow_html=True)
        avg_by_loc = df.groupby("Category")["LST_C"].agg(["mean", "max", "min"]).round(2)
        table_html = generate_html_table(avg_by_loc, st.session_state["theme"] == "dark")
        st.markdown(table_html, unsafe_allow_html=True)

        col_left, col_right = st.columns(2)

        # Plotly layout configuration variables
        is_dark = st.session_state["theme"] == "dark"
        plot_bg = "#181818" if is_dark else "#FFFFFF"
        grid_color = "#2D2D2D" if is_dark else "#D9E0EA"
        text_color = "#F5F5F5" if is_dark else "#111827"
        muted_text_color = "#9A9A9A" if is_dark else "#6B7280"

        with col_left:
            st.markdown('<div class="section-title">Heat Zone Distribution</div>', unsafe_allow_html=True)
            
            zone_order = ["Critical", "High", "Moderate", "Normal"]
            zone_counts = df["Zone"].value_counts().reindex(zone_order, fill_value=0).reset_index()
            zone_counts.columns = ["Zone", "Count"]

            # Standardized accent mapping
            colors_map = {
                "Critical": "#DC2626", # Red
                "High": "#D97706",     # Orange
                "Moderate": "#F59E0B", # Yellow/Amber
                "Normal": "#059669"    # Green
            }

            fig_bar = px.bar(
                zone_counts,
                x="Zone",
                y="Count",
                color="Zone",
                color_discrete_map=colors_map,
                text="Count"
            )
            fig_bar.update_traces(
                textposition="outside",
                texttemplate="%{text}",
                hovertemplate="<b>%{x} Zone</b><br>Count: %{y}<extra></extra>",
                textfont=dict(color=text_color, size=12, family="Poppins")
            )
            fig_bar.update_layout(
                margin=dict(t=40, b=20, l=40, r=20),
                paper_bgcolor="rgba(0,0,0,0)",
                plot_bgcolor="rgba(0,0,0,0)",
                font=dict(color=text_color, size=12, family="Poppins"),
                showlegend=False,
                xaxis=dict(
                    title="",
                    showgrid=False,
                    linecolor=grid_color,
                    tickfont=dict(size=11, color=muted_text_color)
                ),
                yaxis=dict(
                    title=dict(
                        text="Number of Grid Cells",
                        font=dict(size=12, color=muted_text_color)
                    ),
                    showgrid=True,
                    gridcolor=grid_color,
                    linecolor=grid_color,
                    tickfont=dict(size=11, color=muted_text_color)
                )
            )
            st.plotly_chart(fig_bar, use_container_width=True, config={"displayModeBar": False})

        with col_right:
            st.markdown('<div class="section-title">LST Temporal Trend</div>', unsafe_allow_html=True)
            
            pivot = df.pivot_table(values="LST_C", index="Date", columns="Category", aggfunc="mean").reset_index()
            df_melted = pivot.melt(id_vars=["Date"], var_name="Location", value_name="LST")

            location_colors = {
                "Hubballi": "#2563EB",
                "Dharwad": "#10B981",
                "Hubballi-Central": "#EF4444",
                "Dharwad-Central": "#F59E0B"
            }
            
            def hex_to_rgba(hex_str, alpha=0.08):
                hex_str = hex_str.lstrip('#')
                if len(hex_str) == 6:
                    r, g, b = tuple(int(hex_str[i:i+2], 16) for i in (0, 2, 4))
                    return f"rgba({r}, {g}, {b}, {alpha})"
                return f"rgba(128, 128, 128, {alpha})"

            fig_line = px.line(
                df_melted,
                x="Date",
                y="LST",
                color="Location",
                color_discrete_map=location_colors,
                markers=True
            )
            
            fig_line.update_traces(
                fill="tozeroy",
                line=dict(width=3),
                marker=dict(size=8, symbol="circle", line=dict(width=1.5, color=plot_bg))
            )

            for trace in fig_line.data:
                color = trace.line.color
                trace.fillcolor = hex_to_rgba(color, 0.08)

            fig_line.update_layout(
                margin=dict(t=40, b=20, l=40, r=20),
                paper_bgcolor="rgba(0,0,0,0)",
                plot_bgcolor="rgba(0,0,0,0)",
                font=dict(color=text_color, size=12, family="Poppins"),
                legend=dict(
                    orientation="h",
                    yanchor="bottom",
                    y=1.02,
                    xanchor="right",
                    x=1,
                    title=None,
                    font=dict(size=11, color=text_color)
                ),
                xaxis=dict(
                    title="",
                    showgrid=True,
                    gridcolor=grid_color,
                    linecolor=grid_color,
                    tickfont=dict(size=11, color=muted_text_color)
                ),
                yaxis=dict(
                    title=dict(
                        text="Land Surface Temp (°C)",
                        font=dict(size=12, color=muted_text_color)
                    ),
                    showgrid=True,
                    gridcolor=grid_color,
                    linecolor=grid_color,
                    tickfont=dict(size=11, color=muted_text_color)
                )
            )
            st.plotly_chart(fig_line, use_container_width=True, config={"displayModeBar": False})

    with tab3:
        st.markdown('<div class="section-title">AI-Generated Mitigation Report</div>', unsafe_allow_html=True)
        st.markdown('<div class="report-container">', unsafe_allow_html=True)
        if not gemini_key:
            st.markdown("""
<div class="info-banner">
    Enter your <strong>Gemini API key</strong> in the sidebar (Settings section)
    to generate the AI mitigation strategy report.
</div>
""", unsafe_allow_html=True)
        else:
            if st.button("Generate Report", type="primary"):
                with st.spinner("Generating report with Gemini AI..."):
                    report_text = generate_report(gemini_key, ward_name, stats, factors, avg_temps)
                    st.session_state["report"] = report_text

            if "report" in st.session_state:
                st.markdown(st.session_state["report"])
                docx_buf = create_docx(st.session_state["report"], ward_name, stats)
                st.download_button(
                    "Download DOCX Report",
                    data=docx_buf,
                    file_name=f"UHI_Report_{ward_name}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        st.markdown('</div>', unsafe_allow_html=True)

else:
    # Premium Welcome / Empty State Screen
    st.markdown("""
<div class="welcome-screen">
    <div class="welcome-hero">
        <h1 class="welcome-title">Urban Climate Observatory</h1>
        <p class="welcome-subtitle">Advanced Land Surface Temperature (LST) Analysis & Mitigation Planning Engine</p>
    </div>
    
    <div class="welcome-grid">
        <div class="welcome-card">
            <div class="welcome-card-header">
                <span class="welcome-card-tag">Diagnostics</span>
                <h3>ECOSTRESS LST Analysis</h3>
            </div>
            <p class="welcome-card-desc">Processes high-resolution thermal data from NASA's ECOSTRESS sensor to monitor micro-climate patterns and LST trends across municipal wards.</p>
        </div>
        
        <div class="welcome-card">
            <div class="welcome-card-header">
                <span class="welcome-card-tag">Hotspots</span>
                <h3>Statistical Hotspot Mapping</h3>
            </div>
            <p class="welcome-card-desc">Identifies critical spatial anomalies by classifying surface temperatures against statistical standard deviations into Normal, Moderate, High, and Critical zones.</p>
        </div>
        
        <div class="welcome-card">
            <div class="welcome-card-header">
                <span class="welcome-card-tag">Mitigation</span>
                <h3>AI Adaptation Planner</h3>
            </div>
            <p class="welcome-card-desc">Leverages Gemini Generative AI to generate context-specific urban planning proposals, native species recommendations, and prioritized action matrixes.</p>
        </div>
    </div>

    <div class="welcome-cta-card">
        <div class="cta-status-indicator"></div>
        <div class="cta-content">
            <h4>System Ready for Analysis</h4>
            <p>Please upload a valid NASA ECOSTRESS LST dataset (.csv) in the sidebar configuration panel to populate the analytics dashboard.</p>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)