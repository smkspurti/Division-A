"""
utils/exporter.py
Builds the Word (.docx) export with English + Kannada content
and a summary metrics table.
"""

import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd


def _set_cell_bg(cell, hex_color: str):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_heading(doc, text, level=1, color="1B5E20"):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def _add_table(doc, headers: list, rows: list, header_color="2E7D32"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        _set_cell_bg(cell, header_color)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8.5)
            if r_idx % 2 == 0:
                _set_cell_bg(cell, "E8F5E9")

    return table


def build_docx(
    plan_english: str,
    plan_kannada: str,
    pest_list: list[dict],
    schemes: list[dict],
    loss_df: pd.DataFrame,
    crop: str,
    quantity: float,
    moisture: float,
    region: str,
    storage_type: str,
) -> io.BytesIO:
    doc = Document()

    # ── Page margins ───────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Cover block ────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("🌾  Post-Harvest Loss Reduction Advisor")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(27, 94, 32)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Generated: {datetime.now().strftime('%d %B %Y, %I:%M %p')}  |  Crop: {crop}  |  Region: {region}").font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()

    # ── Farmer Summary Table ───────────────────────────────────────────────
    _add_heading(doc, "Farmer Input Summary", level=2)
    _add_table(
        doc,
        headers=["Parameter", "Value"],
        rows=[
            ["Crop Type",      crop],
            ["Harvest Quantity", f"{quantity} kg"],
            ["Current Moisture", f"{moisture}%"],
            ["Region",          region],
            ["Storage Type",    storage_type],
            ["Report Date",     datetime.now().strftime("%d-%m-%Y")],
        ],
    )
    doc.add_paragraph()

    # ── Post-Harvest Loss Data ─────────────────────────────────────────────
    if not loss_df.empty and "loss_pct" in loss_df.columns:
        _add_heading(doc, "Expected Post-Harvest Losses (FAO Data)", level=2)
        has_cause = "cause" in loss_df.columns
        if has_cause:
            rows = [
                (r.get("stage", "Processing"), f"{r['loss_pct']}%", r.get("cause", "—"))
                for _, r in loss_df.iterrows()
            ]
            total_pct = loss_df["loss_pct"].sum()
            total_kg  = round(quantity * total_pct / 100, 1)
            rows.append(["TOTAL", f"{round(total_pct,1)}%", f"≈ {total_kg} kg potential loss"])
            _add_table(doc, ["Stage", "Loss %", "Main Cause"], rows)
        else:
            rows = [
                (r.get("stage", "Processing"), f"{r['loss_pct']}%")
                for _, r in loss_df.iterrows()
            ]
            total_pct = loss_df["loss_pct"].sum()
            total_kg  = round(quantity * total_pct / 100, 1)
            rows.append(["TOTAL", f"{round(total_pct,1)}%"])
            doc.add_paragraph(f"Estimated loss: {total_kg} kg of {crop} at risk without intervention.")
            _add_table(doc, ["Stage", "Loss %"], rows)
        doc.add_paragraph()

    # ── Management Plan (English) ──────────────────────────────────────────
    _add_heading(doc, "Post-Harvest Management Plan (English)", level=1)
    for para_text in plan_english.split("\n\n"):
        if para_text.strip():
            p = doc.add_paragraph(para_text.strip())
            p.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()

    # ── Pest Risk Calendar ─────────────────────────────────────────────────
    if pest_list:
        _add_heading(doc, "Pest Risk Calendar", level=2)
        rows = [
            (p["pest"], p["peak_months"], p["risk"], p["damage"], p["control"])
            for p in pest_list
        ]
        _add_table(
            doc,
            headers=["Pest", "Peak Season", "Risk", "Damage", "Control Method"],
            rows=rows,
        )
        doc.add_paragraph()

    # ── Government Schemes ─────────────────────────────────────────────────
    _add_heading(doc, "Government Scheme Eligibility", level=2)
    if schemes:
        rows = [
            (s["scheme_name"], s["authority"],
             f"{s['subsidy_pct']}%" if s["subsidy_pct"] else "Free storage",
             s["contact"])
            for s in schemes
        ]
        _add_table(doc, ["Scheme", "Authority", "Benefit", "Contact"], rows)
        doc.add_paragraph()
        for s in schemes:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(s["scheme_name"] + ": ").bold = True
            p.add_run(s["notes"])
    else:
        doc.add_paragraph("⚠️ No central schemes matched. Contact your District Agriculture Officer for local support.")

    doc.add_page_break()

    # ── Kannada Section ────────────────────────────────────────────────────
    kn_title = doc.add_paragraph()
    kn_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kn_title.add_run("🌾  ಕೊಯ್ಲು ನಂತರದ ನಷ್ಟ ಕಡಿತ ಸಲಹೆಗಾರ")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(27, 94, 32)

    doc.add_paragraph()
    _add_heading(doc, "ಕನ್ನಡ ಆವೃತ್ತಿ — ನಿರ್ವಹಣಾ ಯೋಜನೆ", level=1)
    for para_text in plan_kannada.split("\n\n"):
        if para_text.strip():
            p = doc.add_paragraph(para_text.strip())
            p.paragraph_format.space_after = Pt(6)

    # ── Footer note ────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(
        "Sources: FAO Food Loss & Waste Database\n"
        "Team A15 — Post-Harvest Loss Reduction Advisor | AI & Rural Technology Hackathon"
    )
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(120, 120, 120)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer