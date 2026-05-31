# modules/doc_generator.py
# Generates a professional Word (.docx) file from the application text

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import date

def generate_docx(application_text: str, data: dict, output_path: str = "outputs/mutation_application.docx") -> str:
    """
    Creates a professional DOCX file.
    Returns the path to the generated file.
    """
    # Create document
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # ── Header: Government Title ──
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("GOVERNMENT OF KARNATAKA")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1a, 0x53, 0x76)

    sub_header = doc.add_paragraph()
    sub_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub_header.add_run("Department of Revenue — Land Mutation Application")
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Divider line
    doc.add_paragraph("─" * 70)

    # ── Application Number & Date ──
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta.add_run(f"Date: {date.today().strftime('%d/%m/%Y')}").italic = True

    doc.add_paragraph()  # spacing

    # ── Addressee ──
    addr = doc.add_paragraph()
    addr.add_run("TO,\n").bold = True
    addr.add_run(f"The Tahsildar,\n")
    addr.add_run(f"{data.get('taluk', '')} Taluk,\n")
    addr.add_run(f"{data.get('district', '')} District, Karnataka.")

    doc.add_paragraph()

    # ── Subject Line ──
    subj = doc.add_paragraph()
    subj.add_run("SUBJECT: ").bold = True
    mutation_type = data.get('mutation_type', 'Sale')
    subj.add_run(f"Application for Land Mutation — {mutation_type} | Survey No. {data.get('survey_no', '')}")

    doc.add_paragraph("─" * 70)

    # ── Salutation ──
    doc.add_paragraph("Respected Sir/Madam,")
    doc.add_paragraph()

    # ── Body ──
    body = doc.add_paragraph()
    body.add_run(
        f"I, {data.get('applicant_name', '')}, hereby submit this application "
        f"for mutation of land records in {data.get('village', '')}, "
        f"{data.get('taluk', '')} Taluk, {data.get('district', '')} District, Karnataka."
    )

    doc.add_paragraph()

    # ── Property Details Table ──
    prop_heading = doc.add_paragraph()
    prop_heading.add_run("PROPERTY DETAILS:").bold = True

    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"

    details = [
        ("Survey Number", data.get("survey_no", "")),
        ("Village", data.get("village", "")),
        ("Taluk", data.get("taluk", "")),
        ("District", data.get("district", "")),
        ("Land Area", f"{data.get('land_area', '')} Acres"),
        ("Mutation Type", mutation_type),
    ]

    for i, (label, value) in enumerate(details):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = str(value)
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # ── Applicant Details Table ──
    app_heading = doc.add_paragraph()
    app_heading.add_run("APPLICANT DETAILS:").bold = True

    table2 = doc.add_table(rows=3, cols=2)
    table2.style = "Table Grid"

    aadhaar = data.get("aadhaar", "000000000000")
    app_details = [
        ("Full Name", data.get("applicant_name", "")),
        ("Aadhaar Number", f"XXXX-XXXX-{aadhaar[-4:]}"),
        ("Mobile Number", data.get("mobile", "")),
    ]

    for i, (label, value) in enumerate(app_details):
        row = table2.rows[i]
        row.cells[0].text = label
        row.cells[1].text = str(value)
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # ── Declaration ──
    decl = doc.add_paragraph()
    decl.add_run(
        "I hereby declare that all the information provided above is true and correct "
        "to the best of my knowledge. All required supporting documents are enclosed "
        "with this application for your kind reference and necessary action."
    )

    doc.add_paragraph()
    doc.add_paragraph("Kindly process this application at the earliest.")
    doc.add_paragraph()
    doc.add_paragraph("Yours faithfully,")
    doc.add_paragraph()
    doc.add_paragraph()

    # ── Signature Section ──
    sig = doc.add_paragraph()
    sig.add_run("________________________\n").bold = True
    sig.add_run(f"{data.get('applicant_name', '')}\n")
    sig.add_run(f"Mobile: {data.get('mobile', '')}\n")
    sig.add_run(f"Date: {date.today().strftime('%d/%m/%Y')}")

    doc.add_paragraph()
    doc.add_paragraph("─" * 70)

    # ── Enclosures ──
    enc_heading = doc.add_paragraph()
    enc_heading.add_run("ENCLOSURES:").bold = True

    from modules.checklist import get_required_documents
    docs = get_required_documents(mutation_type)
    for i, d in enumerate(docs, 1):
        doc.add_paragraph(f"  {i}. {d}")

    # ── Save file ──
    os.makedirs("outputs", exist_ok=True)
    doc.save(output_path)
    return output_path