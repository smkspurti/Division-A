import io
import os
import matplotlib
matplotlib.use('Agg')  # Non-interactive backend
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from typing import Dict, Any, List
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, hex_color: str):
    """Sets background shading of a table cell to the specified hex color."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_table_borders(table):
    """Sets a premium scientific table border style (horizontal rules only, no vertical borders)."""
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    
    # Horizontal borders
    for b_name in ['top', 'bottom', 'insideH']:
        border = OxmlElement(f'w:{b_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # 0.5 pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'B3B3B3')
        tblBorders.append(border)
        
    # Remove vertical borders
    for b_name in ['left', 'right', 'insideV']:
        border = OxmlElement(f'w:{b_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
        
    tblPr.append(tblBorders)

def generate_comparison_chart(dimensions_before: Dict[str, float], dimensions_after: Dict[str, float]) -> io.BytesIO:
    """
    Generates a high-quality grouped bar chart comparing before and after quality dimension scores.
    """
    labels = list(dimensions_before.keys())
    before_vals = [dimensions_before[l] for l in labels]
    after_vals = [dimensions_after[l] for l in labels]
    
    x = np.arange(len(labels))
    width = 0.35
    
    fig, ax = plt.subplots(figsize=(7.5, 4.0), dpi=300)
    
    # Modern professional colors: Crimson (Before), Emerald (After)
    rects1 = ax.bar(x - width/2, before_vals, width, label='Before Cleaning', color='#E06666', edgecolor='none')
    rects2 = ax.bar(x + width/2, after_vals, width, label='After Cleaning (CleanAgent)', color='#2ECC71', edgecolor='none')
    
    ax.set_ylabel('Quality Score (%)', fontsize=10, fontweight='bold', color='#2C3E50')
    ax.set_title('Data Quality Scores: Before vs After Restorations', fontsize=12, fontweight='bold', color='#2C3E50', pad=15)
    ax.set_xticks(x)
    ax.set_xticklabels(labels, fontsize=9, color='#34495E')
    ax.set_ylim(0, 110)
    ax.grid(axis='y', linestyle=':', alpha=0.5, color='#BDC3C7')
    ax.legend(frameon=True, facecolor='#F8F9F9', edgecolor='none', loc='lower right')
    
    # Values labels on top of bars
    def autolabel(rects):
        for rect in rects:
            height = rect.get_height()
            ax.annotate(f'{height:.1f}%',
                        xy=(rect.get_x() + rect.get_width() / 2, height),
                        xytext=(0, 2),  # 2 points vertical offset
                        textcoords="offset points",
                        ha='center', va='bottom', fontsize=7.5, color='#2C3E50')
                        
    autolabel(rects1)
    autolabel(rects2)
    
    # Clean spines
    for spine in ['top', 'right', 'left']:
        ax.spines[spine].set_visible(False)
    ax.spines['bottom'].set_color('#BDC3C7')
    
    fig.tight_layout()
    img_stream = io.BytesIO()
    fig.savefig(img_stream, format='png', bbox_inches='tight')
    img_stream.seek(0)
    plt.close(fig)
    return img_stream

def generate_docx_report(comparison: Dict[str, Any], action_log: List[Dict[str, Any]], initial_issues: List[Dict[str, Any]], output_path: str) -> None:
    """
    Compiles a comprehensive, professional DOCX data quality report
    with embedded charts, methodology details, before-after tables, and action audit trail.
    """
    doc = Document()
    
    # ----------------------------------------------------
    # Styles Setup
    # ----------------------------------------------------
    # Configure margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Setup styles
    styles = doc.styles
    
    # Document Title Style
    title_style = styles.add_style('ReportTitle', 1)
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(24)
    title_font.bold = True
    title_font.color.rgb = RGBColor(44, 62, 80)  # Dark Blue-Gray #2C3E50
    
    # Heading 1 Style
    h1_style = styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Arial'
    h1_font.size = Pt(16)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(41, 128, 185)  # Medium Blue #2980B9
    
    # Heading 2 Style
    h2_style = styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = 'Arial'
    h2_font.size = Pt(13)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(52, 73, 94)
    
    # Normal/Body Text Style
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(51, 51, 51)
    
    # ----------------------------------------------------
    # Cover / Header Title
    # ----------------------------------------------------
    title_p = doc.add_paragraph("DATA QUALITY AUDIT & AUTOMATED RESTORATION REPORT", style='ReportTitle')
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    subtitle_p = doc.add_paragraph("Autonomous Agentic Data Cleaning & 6-Dimension Diagnostics Log")
    subtitle_p.runs[0].font.size = Pt(13)
    subtitle_p.runs[0].font.italic = True
    subtitle_p.runs[0].font.color.rgb = RGBColor(127, 140, 141)
    
    doc.add_paragraph(f"Generated on: {pd.Timestamp.now().strftime('%B %d, %Y at %I:%M %p')}\nSystem: DataCleanAgent Autonomous GenAI System")
    doc.add_paragraph("-" * 80)
    
    # ----------------------------------------------------
    # Section 1: Executive Summary
    # ----------------------------------------------------
    doc.add_heading("1. Executive Summary", level=1)
    
    summary_text = (
        "Data quality issues like missing values, outliers, type conflicts, and labels variations "
        "introduce significant errors in downstream machine learning and analytics workflows. "
        "This report summarizes the diagnostics and automated cleaning actions executed by the "
        "DataCleanAgent autonomous cleaning engine.\n\n"
        f"The raw dataset originally contained {comparison['summary_before']['rows']} records and {comparison['summary_before']['columns']} columns. "
        f"The DataCleanAgent profiling engine scanned the data across 6 primary dimensions (Completeness, Uniqueness, Validity, Consistency, Accuracy, and Integrity). "
        f"Following diagnostics, the autonomous Gemini 2.5 Flash agent executed {len(action_log)} sequenced cleaning steps. "
        f"The overall data quality score increased from {comparison['overall_before']:.2f}% to {comparison['overall_after']:.2f}%, representing a "
        f"net quality improvement of +{comparison['overall_delta']:.2f}%."
    )
    doc.add_paragraph(summary_text)
    
    # Embed Grouped Chart
    dims_before = {dim: val['before'] for dim, val in comparison['dimensions'].items()}
    dims_after = {dim: val['after'] for dim, val in comparison['dimensions'].items()}
    chart_stream = generate_comparison_chart(dims_before, dims_after)
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(chart_stream, width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    caption_p = doc.add_paragraph("Figure 1: Side-by-side comparison of individual data quality dimension scores before and after restoration.")
    caption_p.runs[0].font.size = Pt(9.5)
    caption_p.runs[0].font.italic = True
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ----------------------------------------------------
    # Section 2: Data Quality Methodology
    # ----------------------------------------------------
    doc.add_heading("2. Data Quality Methodology", level=1)
    
    intro_method = (
        "The DataCleanAgent scoring framework is inspired by the ISO 25012 Data Quality Model. "
        "To establish a comprehensive and defensive metric, the Overall Quality Score is calculated "
        "using the weighted harmonic mean of the individual dimension scores. The harmonic mean "
        "strictly penalizes critical dimension-level failures, ensuring that high scores in other areas "
        "cannot mask a complete failure in a critical dimension (e.g. 0% completeness)."
    )
    doc.add_paragraph(intro_method)
    
    # Descriptions of dimensions
    doc.add_heading("Dimension Matrix Definitions", level=2)
    
    dim_defs = [
        ("Completeness", "Measures the presence of values. Scans for standard cell nulls (NaN, None) and resolves textual sentinel nulls ('?', 'N/A', '-', 'none')."),
        ("Uniqueness", "Scans for duplicate rows. Evaluates records after excluding auto-incrementing index columns to isolate true duplication."),
        ("Validity", "Validates values against data type models and logical range bounds (e.g. Survived in [0, 1], Pclass in [1, 2, 3], Age between 0 and 120, non-negative Fares)."),
        ("Consistency", "Standardizes categorical labels. Matches categories case-insensitively and computes pairwise string similarities via the Levenshtein-like SequenceMatcher ratio to cluster and unify abbreviations (e.g. Male/male/M, female/fem/F)."),
        ("Accuracy", "Detects numerical and statistical outliers. Integrates Interquartile Range (IQR) bounds, standard deviations (Z-score > 3), and multivariate Isolation Forest detection to identify anomalies."),
        ("Integrity", "Checks structural and relational constraints (e.g., ensuring Passenger IDs are uniquely assigned and checking logic boundaries like baby passenger travelling with parents).")
    ]
    
    for d_title, d_desc in dim_defs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        r_bold = p.add_run(f"•  {d_title}: ")
        r_bold.bold = True
        r_bold.font.color.rgb = RGBColor(52, 73, 94)
        p.add_run(d_desc)
        
    # ----------------------------------------------------
    # Section 3: Before vs After Summary Metrics Table
    # ----------------------------------------------------
    doc.add_heading("3. Data Quality Dimension Assessment", level=1)
    doc.add_paragraph("The table below details the statistical improvement in each quality dimension:")
    
    # Table creation
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Quality Dimension'
    hdr_cells[1].text = 'Before Score'
    hdr_cells[2].text = 'After Score'
    hdr_cells[3].text = 'Net Change'
    
    # Format headers
    for cell in hdr_cells:
        set_cell_background(cell, '2980B9')  # Blue theme
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
    for dim, delta_info in comparison['dimensions'].items():
        row_cells = table.add_row().cells
        row_cells[0].text = dim
        row_cells[1].text = f"{delta_info['before']:.2f}%"
        row_cells[2].text = f"{delta_info['after']:.2f}%"
        
        delta = delta_info['delta']
        sign = "+" if delta >= 0 else ""
        row_cells[3].text = f"{sign}{delta:.2f}%"
        
        # Color code positive/negative deltas
        if delta > 0:
            row_cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(46, 204, 113)  # Green
        elif delta < 0:
            row_cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(192, 57, 43)  # Red
            
    set_table_borders(table)
    
    # Add spacing
    doc.add_paragraph()
    
    # ----------------------------------------------------
    # Section 4: Initial Diagnostics Log
    # ----------------------------------------------------
    doc.add_heading("4. Detected Issues Log (Pre-Cleaning)", level=1)
    
    if not initial_issues:
        doc.add_paragraph("No significant issues were detected in the raw dataset.")
    else:
        doc.add_paragraph(f"The profiling engine isolated {len(initial_issues)} quality issues in the raw data:")
        
        table_issues = doc.add_table(rows=1, cols=4)
        table_issues.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_i = table_issues.rows[0].cells
        hdr_i[0].text = 'Dimension'
        hdr_i[1].text = 'Column'
        hdr_i[2].text = 'Severity'
        hdr_i[3].text = 'Description'
        
        for cell in hdr_i:
            set_cell_background(cell, '34495E')  # Dark header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            
        for issue in initial_issues:
            row_cells = table_issues.add_row().cells
            row_cells[0].text = issue['dimension']
            row_cells[1].text = issue['column']
            row_cells[2].text = issue['severity']
            row_cells[3].text = issue['description']
            
            # Severity color code
            sev = issue['severity']
            run = row_cells[2].paragraphs[0].runs[0]
            run.font.bold = True
            if sev == 'High':
                run.font.color.rgb = RGBColor(192, 57, 43)  # Dark Red
            elif sev == 'Medium':
                run.font.color.rgb = RGBColor(230, 126, 34)  # Orange
            else:
                run.font.color.rgb = RGBColor(127, 140, 141)  # Gray
                
        set_table_borders(table_issues)
        
    doc.add_page_break()
    
    # ----------------------------------------------------
    # Section 5: Autonomous Cleaning Log (Audit Trail)
    # ----------------------------------------------------
    doc.add_heading("5. Autonomous Cleaning Audit Trail", level=1)
    doc.add_paragraph(
        "Below is the complete sequence of cleaning steps planned and executed "
        "by the autonomous GenAI agent. Every action includes a statistical justification:"
    )
    
    table_audit = doc.add_table(rows=1, cols=4)
    table_audit.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_a = table_audit.rows[0].cells
    hdr_a[0].text = 'Step'
    hdr_a[1].text = 'Action / Justification'
    hdr_a[2].text = 'Generated Pandas Code'
    hdr_a[3].text = 'Status'
    
    for cell in hdr_a:
        set_cell_background(cell, '16A085')  # Teal theme
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
    # Column width rules for table readability
    col_widths = [Inches(0.5), Inches(2.5), Inches(2.7), Inches(0.8)]
    
    for idx, act in enumerate(action_log):
        row_cells = table_audit.add_row().cells
        row_cells[0].text = str(act.get('step', idx + 1))
        
        # Action Justification
        row_cells[1].text = act.get('justification', '')
        
        # Pandas Code (styled as monospace code)
        code_p = row_cells[2].paragraphs[0]
        code_p.text = act.get('code', '')
        if code_p.runs:
            code_p.runs[0].font.name = 'Consolas'
            code_p.runs[0].font.size = Pt(8.5)
            code_p.runs[0].font.color.rgb = RGBColor(44, 62, 80)
            
        # Status
        err = act.get('error')
        if err:
            row_cells[3].text = "FAILED"
            row_cells[3].paragraphs[0].runs[0].font.bold = True
            row_cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(192, 57, 43)  # Red
        else:
            row_cells[3].text = "SUCCESS"
            row_cells[3].paragraphs[0].runs[0].font.bold = True
            row_cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(39, 174, 96)  # Green
            
    # Set widths
    for row in table_audit.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width
            
    set_table_borders(table_audit)
    
    doc.add_paragraph()
    
    # ----------------------------------------------------
    # Section 6: Column-Level Quality Deltas
    # ----------------------------------------------------
    doc.add_heading("6. Column-Level Before vs After Comparison", level=1)
    doc.add_paragraph("Detailed breakdown of missing values and types before and after cleaning:")
    
    table_cols = doc.add_table(rows=1, cols=5)
    table_cols.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_c = table_cols.rows[0].cells
    hdr_c[0].text = 'Column Name'
    hdr_c[1].text = 'Inferred Type'
    hdr_c[2].text = 'Nulls (Before)'
    hdr_c[3].text = 'Nulls (After)'
    hdr_c[4].text = 'Outliers (B → A)'
    
    for cell in hdr_c:
        set_cell_background(cell, '8E44AD')  # Purple theme
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
    for col, delta in comparison['columns'].items():
        row_cells = table_cols.add_row().cells
        row_cells[0].text = col
        row_cells[1].text = delta.get('type_before', '')
        
        null_b = delta.get('null_count_before', 0)
        null_a = delta.get('null_count_after', 0)
        row_cells[2].text = str(null_b)
        row_cells[3].text = str(null_a)
        
        # Outlier counts before -> after
        outliers_b = delta.get('outliers_before', 0)
        outliers_a = delta.get('outliers_after', 0)
        row_cells[4].text = f"{outliers_b} → {outliers_a}" if 'outliers_before' in delta else "-"
        
        # Color reductions in nulls
        if null_b > null_a:
            row_cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(39, 174, 96)  # Green
            row_cells[3].paragraphs[0].runs[0].font.bold = True
            
    set_table_borders(table_cols)
    
    # Save Report
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
