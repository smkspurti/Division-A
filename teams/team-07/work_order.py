from docx import Document
from docx.shared import Pt, RGBColor
from datetime import datetime

def get_work_order_template(fault_type, priority, road_type):
    complexity = 'High' if priority == 'P1' else 'Medium' if priority == 'P2' else 'Low'
    hours = '4' if priority == 'P1' else '2' if priority == 'P2' else '1'
    
    if fault_type == "burnt":
        desc = "Complete replacement of the luminaire module due to confirmed burnout. Inspect internal wiring for heat damage before installing new unit."
        materials = "• LED Luminaire Module (230V)\n• Wire Connectors\n• Heat-shrink Tubing"
        safety = "• Isolate mains power at junction before climbing\n• Use insulated gloves and tools"
        notes = "Ensure heat sink is properly seated when installing the new module."
    elif fault_type == "flickering":
        desc = "Inspect and likely replace the LED driver or ballast causing unstable power delivery. Verify incoming voltage stability."
        materials = "• LED Driver Unit (Compatible model)\n• Multimeter\n• Replacement wiring"
        safety = "• Capacitor in driver may retain charge, discharge safely\n• Secure bucket truck in stable position"
        notes = "If incoming voltage fluctuates widely, do not replace driver; report back for grid-level investigation."
    elif fault_type == "voltage_surge":
        desc = "Inspect for surge damage. Replace Surge Protection Device (SPD) and test transformer output for over-voltage."
        materials = "• Surge Protection Device (10kV or higher)\n• Multimeter\n• Standard tool kit"
        safety = "• Potential live wires due to melted insulation, proceed with extreme caution\n• Full PPE required"
        notes = "Check adjacent nodes for similar damage, as surges often affect a localized cluster."
    elif fault_type == "offline":
        desc = "Investigate total loss of power. Check MCB, fuse, and incoming power cable for physical damage or disconnection."
        materials = "• Replacement Fuses/MCB\n• Cable fault locator\n• Splicing kit"
        safety = "• Treat all lines as live until tested\n• Check for water ingress in base compartment"
        notes = "If pole is physically damaged (e.g., from a vehicle collision), secure the area and request civil works team."
    else:
        desc = "General inspection and repair based on on-site findings."
        materials = "• Standard maintenance kit"
        safety = "• Standard electrical safety protocols"
        notes = "Determine root cause and update maintenance log."

    template = (
        f"1. WORK DESCRIPTION\n{desc}\n\n"
        f"2. MATERIALS REQUIRED\n{materials}\n\n"
        f"3. SAFETY PRECAUTIONS\n{safety}\n\n"
        f"4. ESTIMATED TIME: {hours} hours (Complexity: {complexity})\n\n"
        f"5. TECHNICIAN NOTES\n{notes}"
    )
    return template

def generate_work_order(fault_dict, report_text, output_path):
    doc = Document()
    
    doc.add_heading("Street Light Maintenance Work Order", level=0)
    
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M")
    date_str = datetime.now().strftime("%Y%m%d")
    node_id = fault_dict['node_id']
    
    table = doc.add_table(rows=6, cols=2, style="Table Grid")
    table.cell(0,0).text = "Work Order #"
    table.cell(0,1).text = f"WO-{node_id}-{date_str}"
    
    table.cell(1,0).text = "Node ID"
    table.cell(1,1).text = str(node_id)
    
    table.cell(2,0).text = "Road Type"
    table.cell(2,1).text = str(fault_dict['road_type'])
    
    table.cell(3,0).text = "Generated"
    table.cell(3,1).text = now_str
    
    table.cell(4,0).text = "Priority"
    priority = fault_dict['priority']
    p_run = table.cell(4,1).paragraphs[0].add_run(priority)
    p_run.bold = True
    if priority == "P1":
        p_run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    elif priority == "P2":
        p_run.font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)
    elif priority == "P3":
        p_run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
        
    table.cell(5,0).text = "GPS Location"
    table.cell(5,1).text = f"{fault_dict['lat']:.5f}, {fault_dict['lon']:.5f}"
    
    doc.add_heading("Fault Details", level=1)
    fault_type = fault_dict['fault_type']
    doc.add_paragraph(f"Fault Type: {fault_type.upper()}")
    doc.add_paragraph(f"Voltage: {fault_dict['avg_v']:.1f}V | Current: {fault_dict['avg_c']:.3f}A | Power: {fault_dict['avg_p']:.1f}W")
    doc.add_paragraph(f"Brightness: {fault_dict['avg_brightness']:.1f}% | On-ratio: {fault_dict['on_ratio']*100:.1f}%")
    
    doc.add_heading("AI-Generated Fault Report", level=1)
    doc.add_paragraph(report_text)
    
    doc.add_heading("Work Instructions", level=1)
    work_order_text = get_work_order_template(fault_type, priority, fault_dict['road_type'])
    doc.add_paragraph(work_order_text)
    
    doc.add_heading("Standard Maintenance Checklist", level=1)
    checklists = {
        "burnt": ["☐ Isolate power supply", "☐ Replace bulb/LED module", "☐ Test circuit continuity", "☐ Update maintenance log"],
        "flickering": ["☐ Check ballast/driver unit", "☐ Inspect wiring connections", "☐ Test voltage stability", "☐ Replace driver if faulty"],
        "voltage_surge": ["☐ Inspect surge protector", "☐ Check transformer output", "☐ Log voltage readings", "☐ Install/replace SPD"],
        "offline": ["☐ Verify power supply at junction", "☐ Inspect fuse and MCB", "☐ Check cable for physical damage", "☐ Restore and test"],
    }
    checklist = checklists.get(fault_type, ["☐ Inspect unit", "☐ Report findings"])
    for item in checklist:
        doc.add_paragraph(item)
        
    doc.add_paragraph("\nTechnician Sign-off: ________________  Date: ________")
    doc.save(output_path)

def generate_summary_report(faults_df, all_nodes_count, output_path="work_orders/daily_fault_summary.docx"):
    doc = Document()
    doc.add_heading("Daily Street Light Fault Summary Report", level=0)
    
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M")
    today = datetime.now().strftime("%Y-%m-%d")
    
    doc.add_paragraph(f"Generated: {now_str}")
    doc.add_paragraph("Municipality: Bengaluru Smart City Division")
    
    doc.add_heading("Overview", level=1)
    table_over = doc.add_table(rows=5, cols=2, style="Table Grid")
    table_over.cell(0,0).text = "Total Nodes Monitored"
    table_over.cell(0,1).text = str(all_nodes_count)
    
    table_over.cell(1,0).text = "Healthy Nodes"
    table_over.cell(1,1).text = str(all_nodes_count - len(faults_df))
    
    table_over.cell(2,0).text = "Faulty Nodes"
    table_over.cell(2,1).text = str(len(faults_df))
    
    table_over.cell(3,0).text = "Report Date"
    table_over.cell(3,1).text = today
    
    table_over.cell(4,0).text = "Next Review"
    table_over.cell(4,1).text = "Next Day"
    
    doc.add_heading("Fault Breakdown", level=1)
    fault_counts = faults_df["fault_type"].value_counts()
    table_break = doc.add_table(rows=1+len(fault_counts), cols=3, style="Table Grid")
    table_break.cell(0,0).text = "Fault Type"
    table_break.cell(0,1).text = "Count"
    table_break.cell(0,2).text = "% of Total"
    
    total_faults = len(faults_df)
    for i, (ftype, count) in enumerate(fault_counts.items()):
        table_break.cell(i+1, 0).text = str(ftype)
        table_break.cell(i+1, 1).text = str(count)
        table_break.cell(i+1, 2).text = f"{(count/total_faults)*100:.1f}%" if total_faults > 0 else "0%"
        
    doc.add_heading("Priority Summary", level=1)
    priority_counts = faults_df["priority"].value_counts()
    table_pri = doc.add_table(rows=4, cols=3, style="Table Grid")
    table_pri.cell(0,0).text = "Priority"
    table_pri.cell(0,1).text = "Count"
    table_pri.cell(0,2).text = "Action Required"
    
    actions = {
        "P1": "Immediate dispatch (< 2 hours)",
        "P2": "Same day repair",
        "P3": "Schedule within 48 hours"
    }
    
    for i, p in enumerate(["P1", "P2", "P3"]):
        table_pri.cell(i+1, 0).text = p
        table_pri.cell(i+1, 1).text = str(priority_counts.get(p, 0))
        table_pri.cell(i+1, 2).text = actions[p]
        
    doc.add_heading("High Priority Nodes (P1)", level=1)
    p1_df = faults_df[faults_df['priority'] == 'P1']
    table_p1 = doc.add_table(rows=1+len(p1_df), cols=4, style="Table Grid")
    table_p1.cell(0,0).text = "Node ID"
    table_p1.cell(0,1).text = "Road Type"
    table_p1.cell(0,2).text = "Fault Type"
    table_p1.cell(0,3).text = "GPS"
    
    for i, row in enumerate(p1_df.itertuples()):
        table_p1.cell(i+1, 0).text = str(row.node_id)
        table_p1.cell(i+1, 1).text = str(row.road_type)
        table_p1.cell(i+1, 2).text = str(row.fault_type)
        table_p1.cell(i+1, 3).text = f"{row.lat:.5f}, {row.lon:.5f}"
        
    doc.add_heading("Recommendations", level=1)
    doc.add_paragraph("1. Dispatch emergency crews immediately to all P1 Highway nodes.")
    doc.add_paragraph("2. Schedule maintenance teams for P2 Main Road nodes within 8 hours.")
    doc.add_paragraph("3. Queue P3 Residential nodes for routine maintenance cycle.")
    
    doc.save(output_path)
